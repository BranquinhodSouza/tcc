"""
Gerador do Pré-Projeto de TCC em formato DOCX
Tema: Análise de Evasão no PROEJA do IFG Campus Goiânia

Para executar:
  pip install python-docx
  python gerar_pre_projeto.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ========== ESTILOS ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(14)
        hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        hs.font.size = Pt(12)
    else:
        hs.font.size = Pt(12)

# ========== CAPA ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INSTITUTO FEDERAL DE EDUCAÇÃO, CIÊNCIA E TECNOLOGIA DE GOIÁS\n')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('CAMPUS GOIÂNIA\n')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('DEPARTAMENTO DAS ÁREAS ACADÊMICAS IV\n')
run.font.size = Pt(11)
run = p.add_run('COORDENAÇÃO DE INFORMÁTICA\n')
run.font.size = Pt(11)
run = p.add_run('BACHARELADO EM SISTEMAS DE INFORMAÇÃO\n\n')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('ANÁLISE DE EVASÃO NO PROEJA DO IFG CAMPUS GOIÂNIA:\nIDENTIFICAÇÃO DE PADRÕES, PROPOSTAS DE INTERVENÇÃO\nE DESENVOLVIMENTO DE UM DASHBOARD PARA MONITORAMENTO')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Paulo Branquinho de Souza')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Goiânia/GO\n2026')

doc.add_page_break()

# ========== FICHA CATALOGRÁFICA ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Paulo Branquinho de Souza')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run('ANÁLISE DE EVASÃO NO PROEJA DO IFG CAMPUS GOIÂNIA: IDENTIFICAÇÃO DE PADRÕES, PROPOSTAS DE INTERVENÇÃO E DESENVOLVIMENTO DE UM DASHBOARD PARA MONITORAMENTO')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run('Pré-projeto de Trabalho de Conclusão de Curso apresentado no Instituto Federal de Goiás como requisito básico para a conclusão do curso de Sistemas de Informação.')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Orientador: Prof. <título>. <nome>')

doc.add_page_break()

# ========== SUMÁRIO ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUMÁRIO')
run.bold = True
run.font.size = Pt(14)

sumario = [
    ('1 INTRODUÇÃO', 4),
    ('2 JUSTIFICATIVA', 5),
    ('3 OBJETIVOS', 6),
    ('3.1 Objetivo Geral', 6),
    ('3.2 Objetivos Específicos', 6),
    ('4 REFERENCIAL TEÓRICO', 7),
    ('5 METODOLOGIA DA PESQUISA', 9),
    ('5.1 Classificação da Pesquisa', 9),
    ('5.2 Procedimentos Metodológicos', 9),
    ('6 CRONOGRAMA', 11),
    ('REFERÊNCIAS BIBLIOGRÁFICAS', 12),
]
for item, page in sumario:
    p = doc.add_paragraph()
    p.add_run(f'{item} .................................................................... {page}')

doc.add_page_break()

# ========== 1 INTRODUÇÃO ==========
doc.add_heading('1 INTRODUÇÃO', level=1)

intro_text = (
    "A Educação de Jovens e Adultos (EJA) constitui uma modalidade de ensino fundamental para a "
    "promoção da inclusão social e do resgate da cidadania de sujeitos historicamente marginalizados "
    "do processo educacional brasileiro. Conforme dados do Anuário Brasileiro da Educação Básica "
    "(2025), o acesso à EJA caiu 34,5% na última década, registrando 2,4 milhões de matrículas em "
    "2024, o menor número da série histórica. Paralelamente, 29% dos brasileiros de 15 a 64 anos "
    "são analfabetos funcionais, evidenciando a persistência de um grave déficit educacional."
)
doc.add_paragraph(intro_text)

intro_text2 = (
    "No âmbito da Rede Federal de Educação Profissional e Tecnológica, o Programa de Integração da "
    "Educação Profissional com a Educação Básica na Modalidade de Educação de Jovens e Adultos "
    "(PROEJA) foi instituído pelo Decreto nº 5.478/2005, posteriormente ampliado pelo Decreto nº "
    "5.840/2006. No Instituto Federal de Goiás (IFG) - Campus Goiânia, o PROEJA teve início com os "
    "Cursos Técnicos Integrados em Serviços de Alimentação e Cozinha, enfrentando desde sua origem "
    "desafios significativos relacionados à evasão escolar."
)
doc.add_paragraph(intro_text2)

intro_text3 = (
    "Estudos específicos sobre o PROEJA no IFG Campus Goiânia, como os de Pereira (2011) e Castro "
    "e Vitorette (2016), revelam um percurso contraditório na construção do direito à educação. "
    "Apenas a Coordenação de Turismo e Hospitalidade mostrou-se favorável à oferta inicial do "
    "programa, evidenciando resistência institucional. A evasão elevada nos cursos ofertados "
    "demonstra a necessidade de compreender os fatores que levam ao abandono escolar nesta "
    "modalidade."
)
doc.add_paragraph(intro_text3)

intro_text4 = (
    "Diante deste cenário, identifica-se a necessidade de investigar de forma sistemática os "
    "padrões de evasão no PROEJA do IFG Campus Goiânia, integrando análise de dados institucionais "
    "e referencial teórico sobre educação de jovens e adultos. O desafio reside em como a tecnologia "
    "da informação, por meio de ferramentas de Business Intelligence (BI) e análise preditiva, pode "
    "auxiliar na identificação precoce de alunos em risco de abandono e no monitoramento contínuo "
    "dos indicadores de permanência."
)
doc.add_paragraph(intro_text4)

intro_text5 = (
    "Nesse contexto, assume-se que o desenvolvimento de um dashboard interativo para visualização "
    "e análise dos dados de evasão, combinado com a identificação de padrões sociodemográficos e "
    "institucionais, permitirá à gestão acadêmica tomar decisões mais informadas para a elaboração "
    "de políticas de permanência estudantil. Espera-se que esta pesquisa contribua para a redução "
    "dos índices de abandono e para o fortalecimento do PROEJA como política pública de inclusão "
    "educacional."
)
doc.add_paragraph(intro_text5)

# ========== 2 JUSTIFICATIVA ==========
doc.add_heading('2 JUSTIFICATIVA', level=1)

just_text = (
    "A escolha deste tema justifica-se pela relevância social e acadêmica da EJA no Brasil, "
    "modalidade que atende majoritariamente alunos pretos e pardos (76,8% das matrículas), "
    "trabalhadores de baixa renda e pessoas em situação de vulnerabilidade social (Todos pela "
    "Educação, 2025). A evasão escolar no PROEJA representa não apenas o desperdício de recursos "
    "públicos, mas também a perpetuação de um ciclo de exclusão educacional que compromete a "
    "mobilidade social desses indivíduos."
)
doc.add_paragraph(just_text)

just_text2 = (
    "O presente trabalho contribui diretamente para a formação esperada no Perfil do Egresso do "
    "curso de Sistemas de Informação, uma vez que exige a aplicação integrada de conhecimentos "
    "multidisciplinares das seguintes áreas:"
)
doc.add_paragraph(just_text2)

areas = [
    "Banco de Dados: Através da modelagem e estruturação dos dados institucionais de matrícula, "
    "frequência e evasão, aplicando técnicas de modelagem dimensional para suporte à decisão.",
    "Business Intelligence: Com o desenvolvimento de um dashboard interativo utilizando ferramentas "
    "como Power BI ou Google Looker Studio, competência essencial para o profissional de TI "
    "moderno na análise de indicadores educacionais.",
    "Mineração de Dados: Pela aplicação de técnicas de análise exploratória e identificação de "
    "padrões nos dados de evasão, permitindo a construção de modelos preditivos de risco.",
    "Engenharia de Software: Através do levantamento de requisitos e da modelagem do sistema de "
    "monitoramento, garantindo que a solução siga padrões de qualidade e usabilidade.",
]
for area in areas:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(area)

just_text3 = (
    "Do ponto de vista institucional, a pesquisa oferece ao IFG Campus Goiânia um instrumento "
    "concreto para o monitoramento da evasão, permitindo a identificação precoce de alunos em "
    "risco e a avaliação do impacto das políticas de assistência estudantil. Conforme Oliveira "
    "e Carmo (2021), a temática da evasão no PROEJA ainda carece de estudos aprofundados que "
    "integrem análise quantitativa de dados e propostas de intervenção baseadas em evidências."
)
doc.add_paragraph(just_text3)

just_text4 = (
    "A inovação deste trabalho reside na integração de técnicas de análise de dados educacionais "
    "com a proposição de medidas práticas de intervenção, consolidando o papel do Bacharel em "
    "Sistemas de Informação como um agente capaz de projetar, implementar e gerenciar soluções "
    "tecnológicas que impactam positivamente a sociedade e a gestão educacional pública."
)
doc.add_paragraph(just_text4)

# ========== 3 OBJETIVOS ==========
doc.add_heading('3 OBJETIVOS', level=1)

doc.add_heading('3.1 Objetivo Geral', level=2)
doc.add_paragraph(
    "Analisar os fatores associados à evasão no PROEJA do IFG Campus Goiânia, identificando "
    "padrões sociodemográficos e institucionais, propondo medidas de intervenção e "
    "desenvolvendo um dashboard interativo para monitoramento contínuo dos indicadores de "
    "permanência estudantil."
)

doc.add_heading('3.2 Objetivos Específicos', level=2)
especificos = [
    "Levantar o referencial teórico sobre evasão escolar na EJA e no PROEJA, com ênfase em "
    "estudos realizados no IFG Campus Goiânia;",
    "Identificar os principais fatores que contribuem para a evasão nos cursos PROEJA do IFG "
    "Campus Goiânia, a partir da análise de dados institucionais e da literatura existente;",
    "Mapear o perfil sociodemográfico dos alunos evadidos (faixa etária, sexo, raça, renda, "
    "situação de trabalho, presença de filhos) para identificação de grupos de risco;",
    "Elaborar um conjunto de propostas de intervenção institucional, pedagógica e de apoio ao "
    "estudante visando a redução das taxas de evasão;",
    "Modelar e desenvolver um dashboard interativo para visualização dos indicadores de evasão, "
    "utilizando ferramentas de Business Intelligence;",
    "Avaliar a eficácia do dashboard como ferramenta de apoio à gestão acadêmica na tomada de "
    "decisões relacionadas à permanência estudantil.",
]
for i, obj in enumerate(especificos, 1):
    doc.add_paragraph(f"{i}. {obj}")

# ========== 4 REFERENCIAL TEÓRICO ==========
doc.add_heading('4 REFERENCIAL TEÓRICO', level=1)

ref1 = (
    "A fundamentação teórica deste projeto sustenta-se na necessidade de compreender a evasão "
    "escolar na Educação de Jovens e Adultos como um fenômeno multidimensional, fortemente "
    "influenciado por fatores socioeconômicos, institucionais e pedagógicos. Nessa perspectiva, "
    "o trabalho ancora-se inicialmente nos estudos de Oliveira e Carmo (2021), que em sua "
    "revisão integrativa sobre evasão no PROEJA identificaram duas grandes categorias de causas: "
    "a necessidade de trabalho/remuneração e a falta de preparo dos professores para atuar na "
    "modalidade."
)
doc.add_paragraph(ref1)

ref2 = (
    "Para compreender o contexto específico do IFG Campus Goiânia, a pesquisa fundamenta-se nos "
    "trabalhos de Pereira (2011), que investigou os fatores de acesso e permanência no PROEJA, "
    "e de Castro e Vitorette (2016), que analisaram as contradições presentes no processo de "
    "implantação e implementação do programa no campus. Esses estudos revelam que a resistência "
    "institucional, a focalização do programa e a inadequação das metodologias de ensino "
    "contribuem significativamente para o abandono escolar."
)
doc.add_paragraph(ref2)

ref3 = (
    "Do ponto de vista das políticas públicas, o trabalho dialoga com os dados do Anuário "
    "Brasileiro da Educação Básica (2025) e com a análise de Silva, Silva e Oliveira (2023) "
    "sobre o cumprimento das metas 8, 9 e 10 do Plano Nacional de Educação (PNE 2014-2024) "
    "em Goiás e Goiânia. As autoras constatam que as metas estão distantes de se concretizar, "
    "com descontinuidade de estratégias e regressão dos indicadores nos últimos anos."
)
doc.add_paragraph(ref3)

ref4 = (
    "No campo da análise de dados educacionais, o referencial teórico expande-se para as "
    "técnicas de Business Intelligence (BI) aplicadas à educação. Conforme destacado por "
    "Assis (2017), a mineração de dados educacionais permite identificar perfis de estudantes "
    "com propensão à evasão, utilizando algoritmos de classificação como CART, Random Forest "
    "e Regressão Logística. Essas técnicas podem ser combinadas com dashboards interativos "
    "para fornecer à gestão acadêmica informações em tempo real sobre o risco de abandono."
)
doc.add_paragraph(ref4)

ref5 = (
    "Complementarmente, a abordagem de Freire (2023) sobre uma educação dialógica e "
    "emancipadora serve como contraponto às práticas pedagógicas tradicionais identificadas "
    "como inadequadas para a EJA. A valorização dos saberes prévios dos educandos e a "
    "contextualização do currículo são apontadas por Dutra, Melo e Silva (2026) como "
    "estratégias fundamentais para promover o engajamento e a permanência dos estudantes "
    "na modalidade."
)
doc.add_paragraph(ref5)

ref_list = [
    "REVISÃO BIBLIOGRÁFICA sobre o tema da evasão na EJA e PROEJA nos artigos, dissertações "
    "e teses mapeados na pesquisa;",
    "ANÁLISE DOCUMENTAL do Projeto Político-Pedagógico dos cursos PROEJA do IFG Campus Goiânia "
    "e dos relatórios institucionais de evasão;",
    "COLETA DE DADOS quantitativos junto à Coordenação de Registro Acadêmico do IFG Campus "
    "Goiânia, incluindo informações sociodemográficas, de matrícula e de evasão dos alunos;",
    "APLICAÇÃO de técnicas de análise estatística descritiva e preditiva para identificação "
    "de padrões de evasão;",
    "DESENVOLVIMENTO de dashboard em ferramenta de BI (Power BI ou Looker Studio) para "
    "visualização interativa dos indicadores.",
]
doc.add_paragraph(
    "Todas as teorias e dados apresentados servirão de base para a construção do dashboard "
    "e para a interpretação dos padrões de evasão identificados, permitindo a elaboração de "
    "recomendações baseadas em evidências para a gestão acadêmica do IFG Campus Goiânia."
)

doc.add_paragraph(
    "A construção deste trabalho é uma síntese integradora de diversas competências adquiridas "
    "ao longo do Bacharelado em Sistemas de Informação. As disciplinas abaixo contribuem de "
    "forma predominante:"
)

disc = [
    "Banco de Dados: Essencial para a modelagem e estruturação dos dados institucionais, "
    "aplicando conceitos de modelagem dimensional e otimização de consultas para suporte à decisão.",
    "Business Intelligence: Fundamental para o desenvolvimento do dashboard e a aplicação de "
    "técnicas de visualização de dados e indicadores de desempenho (KPIs).",
    "Engenharia de Software: Norteará a fase de levantamento de requisitos e modelagem do "
    "sistema de monitoramento, garantindo uma arquitetura robusta e documentada.",
    "Interação Homem-Computador (IHC): Guiará o desenvolvimento da interface do dashboard, "
    "focando na usabilidade e na experiência do usuário gestor.",
    "Probabilidade e Estatística: Base para a análise descritiva dos dados e para a "
    "aplicação de técnicas de correlação e regressão.",
]
for d in disc:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d)

# ========== 5 METODOLOGIA ==========
doc.add_heading('5 METODOLOGIA DA PESQUISA', level=1)

doc.add_heading('5.1 Classificação da Pesquisa', level=2)
doc.add_paragraph(
    "Quanto à sua natureza, esta é uma pesquisa aplicada, pois objetiva gerar conhecimentos "
    "para aplicação prática na solução de um problema real de evasão escolar no IFG Campus "
    "Goiânia. Do ponto de vista dos objetivos, a pesquisa é exploratória e descritiva, uma "
    "vez que envolve a caracterização detalhada do fenômeno da evasão e a identificação de "
    "padrões. Quanto à abordagem, caracteriza-se como mista (quantitativa e qualitativa): "
    "quantitativa na análise estatística dos dados de evasão e na validação dos indicadores "
    "do dashboard; qualitativa na interpretação dos fatores contextuais e institucionais."
)

doc.add_heading('5.2 Procedimentos Metodológicos', level=2)
doc.add_paragraph(
    "O desenvolvimento será dividido em seis fases fundamentais, seguindo as melhores "
    "práticas de pesquisa científica e engenharia de software:"
)

fases = [
    ("Revisão Sistemática da Literatura",
     "Realização de levantamento bibliográfico em bases como Google Acadêmico, SciELO, "
     "Portal de Periódicos CAPES e Repositório da UnB, focando em artigos, dissertações e "
     "teses sobre evasão na EJA e no PROEJA, com ênfase em estudos sobre o IFG Campus Goiânia."),
    ("Coleta e Preparação dos Dados",
     "Solicitação de dados institucionais junto à Coordenação de Registro Acadêmico do IFG "
     "Campus Goiânia. Os dados incluirão registros de matrícula, histórico acadêmico, "
     "frequência e registros de evasão dos alunos dos cursos PROEJA. Os dados serão "
     "anonimizados para garantir a privacidade dos estudantes, conforme a Lei Geral de "
     "Proteção de Dados (LGPD)."),
    ("Análise Exploratória de Dados (EDA)",
     "Aplicação de técnicas de análise estatística descritiva utilizando Python (pandas, "
     "numpy, scipy) para caracterização do perfil dos alunos evadidos, identificação de "
     "correlações entre variáveis sociodemográficas e evasão, e segmentação por grupos de "
     "risco."),
    ("Modelagem Preditiva",
     "Desenvolvimento de modelos de classificação (Regressão Logística, Random Forest) "
     "para identificação precoce de alunos com propensão à evasão, utilizando variáveis "
     "como frequência, desempenho acadêmico, turno, idade, sexo e situação de trabalho."),
    ("Desenvolvimento do Dashboard",
     "Criação de dashboard interativo utilizando Power BI ou Google Looker Studio, com "
     "painéis de indicadores (KPIs), gráficos de evolução temporal, segmentação por curso "
     "e perfil, e tabela de alerta precoce para alunos em risco."),
    ("Elaboração de Propostas de Intervenção",
     "A partir dos padrões identificados e da revisão da literatura, será elaborado um "
     "conjunto de recomendações institucionais, pedagógicas e de apoio ao estudante, "
     "organizadas em ordem de prioridade e viabilidade de implementação."),
]

for titulo, desc in fases:
    p = doc.add_paragraph()
    run = p.add_run(f"{titulo}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "A validação do dashboard será realizada por meio de testes de usabilidade com gestores "
    "acadêmicos do IFG Campus Goiânia e pela avaliação da acurácia dos modelos preditivos "
    "utilizando métricas como Precisão, Revocação e Curva ROC."
)

# ========== 6 CRONOGRAMA ==========
doc.add_heading('6 CRONOGRAMA', level=1)

table = doc.add_table(rows=9, cols=9)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Atividade / Mês', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

cron = [
    ['1. Revisão Bibliográfica e Refinamento do Tema', 'X', 'X', '', '', '', '', '', ''],
    ['2. Coleta e Preparação dos Dados', '', 'X', 'X', '', '', '', '', ''],
    ['3. Escrita do Relatório de Qualificação (TCC 1)', 'X', 'X', 'X', 'X', '', '', '', ''],
    ['4. Análise Exploratória e Modelagem Preditiva', '', '', '', 'X', 'X', '', '', ''],
    ['5. Desenvolvimento do Dashboard', '', '', '', '', 'X', 'X', '', ''],
    ['6. Elaboração das Propostas de Intervenção', '', '', '', '', '', 'X', '', ''],
    ['7. Escrita da Monografia Final (TCC 2)', '', '', '', '', '', 'X', 'X', 'X'],
    ['8. Preparação para Defesa', '', '', '', '', '', '', '', 'X'],
]

for i, row_data in enumerate(cron):
    for j, val in enumerate(row_data):
        cell = table.rows[i + 1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
        # First column left-aligned
        if j == 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ========== REFERÊNCIAS ==========
doc.add_page_break()
doc.add_heading('REFERÊNCIAS BIBLIOGRÁFICAS', level=1)

refs = [
    "ANUÁRIO BRASILEIRO DA EDUCAÇÃO BÁSICA 2025. São Paulo: Todos pela Educação, 2025.",
    "ASSIS, L. R. S. Perfil de evasão no ensino superior brasileiro: uma abordagem de mineração de dados. 2017. Dissertação (Mestrado) — Universidade de Brasília, Brasília, 2017.",
    "CASTRO, M. D. R.; VITORETTE, J. M. B. O Programa de Integração da Educação Profissional com a Educação Básica na Modalidade de Educação de Jovens e Adultos (PROEJA) no IFG - Câmpus Goiânia: um percurso contraditório na construção do direito à educação. HOLOS, v. 2, p. 301-318, 2016.",
    "CASTRO, M. D. R. O PROEJA no IFG - Câmpus Goiânia: contradições, limites e perspectivas. 2018. Dissertação (Mestrado em Educação) — PUC Goiás, Goiânia, 2018.",
    "DUTRA, M. F. C.; MELO, A. C. L.; SILVA, R. I. P. Desafios e Perspectivas da Educação de Jovens e Adultos (EJA): um olhar sobre a evasão e a prática pedagógica no Brasil Contemporâneo. REASE, v. 12, n. 5, 2026.",
    "FREIRE, P. Pedagogia do Oprimido. 87. ed. Rio de Janeiro: Paz e Terra, 2023.",
    "NERY, M. C. R.; PEREIRA, J. O.; AMARAL, P. G. Educação de Jovens e Adultos - EJA - seus impasses e suas contradições. Revista Eletrônica Científica da UERGS, v. 6, n. 1, p. 29-41, 2020.",
    "OLIVEIRA, P. L.; CARMO, N. C. A temática evasão escolar no contexto do PROEJA: uma revisão integrativa. Revista Ponto de Vista, v. 10, n. 1, p. 01-21, 2021.",
    "PEREIRA, J. V. O PROEJA no Instituto Federal de Goiás – Campus Goiânia: um estudo sobre os fatores de acesso e permanência na escola. 2011. 154 f. Dissertação (Mestrado em Educação) — Universidade de Brasília, Brasília, 2011.",
    "SANTOS, Á. F. et al. Desafios da Educação de Jovens e Adultos no Brasil: entre metodologias inadequadas e políticas instáveis. Missioneira, v. 27, n. 1, p. 103-114, 2025.",
    "SILVA, T. G. C.; SILVA, R. B. B.; OLIVEIRA, G. L. Educação de Jovens e Adultos em Goiás e Goiânia à luz dos dados do INEP. Revista Sapiência, v. 12, n. 2, p. 88-104, 2023.",
    "TOLENTINO FILHO, D. Educação de jovens e adultos no Brasil: Avanços, desafios e perspectivas. International Integralize Scientific, v. 5, n. 46, 2025.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(ref)
    run.font.size = Pt(11)

# ========== SALVAR ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Pre_projeto_Evasao_PROEJA_IFG.docx')
doc.save(output_path)
print(f"Documento salvo em: {output_path}")
print("Pré-projeto gerado com sucesso!")
